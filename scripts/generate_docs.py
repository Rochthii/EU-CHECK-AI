import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Ensure output directories exist
os.makedirs("docs", exist_ok=True)
os.makedirs("Content", exist_ok=True)

# Color Palette Constants (Themis Regal Theme)
COLOR_EMERALD = RGBColor(11, 59, 36)      # #0B3B24
COLOR_GOLD = RGBColor(212, 175, 55)       # #D4AF37
COLOR_DARK = RGBColor(8, 11, 9)          # #080B09
COLOR_GRAY = RGBColor(109, 115, 104)     # #6D7368
HEX_EMERALD = "0B3B24"
HEX_GOLD = "D4AF37"
HEX_LIGHT_BG = "F4F6F4"
HEX_BORDER = "D4AF37"

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{margin_name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def add_header_footer(doc, doc_title):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"COFFEE EU-CHECK AI  |  {doc_title}  |  THEMIS REGAL LEGAL-TECH"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.style.font.name = "Arial"
    hp.style.font.size = Pt(8.5)
    hp.style.font.color.rgb = COLOR_GRAY

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "CHỨNG NHẬN TIỀN KIỂM BẢO MẬT 5 NĂM (EUDR ART. 31) — KHÔNG THAY THẾ CHO BẢO MẬT HẢI QUAN BẮC BỘ"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.style.font.name = "Arial"
    fp.style.font.size = Pt(8.0)
    fp.style.font.color.rgb = COLOR_GRAY

def create_styled_doc(title_text, subtitle_text):
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_header_footer(doc, title_text.upper())

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(title_text)
    run_title.font.name = "Georgia"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_EMERALD

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(subtitle_text)
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_GOLD

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    return doc

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_EMERALD
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_GOLD
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_EMERALD
    return p

def add_body_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    r_body = p.add_run(text)
    r_body.font.name = "Arial"
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = COLOR_DARK
    return p

# ==============================================================================
# 1. GENERATE FEATURE DOC
# ==============================================================================
print("Generating Feature Doc...")

doc_feat = create_styled_doc(
    "TÀI LIỆU ĐỘNG CƠ TÍNH NĂNG TIỀN KIỂM (FEATURE DOC)",
    "Coffee EU-Check AI — Themis Regal Legal-Tech Pre-Audit System for EUDR & CBAM"
)

add_heading_1(doc_feat, "MỤC LỤC TỰ ĐỘNG (TABLE OF CONTENTS)")
add_body_p(doc_feat, "1. Tổng Quan Kiến Trúc Động Cơ Kép (Dual GIS & AI Engine)")
add_body_p(doc_feat, "2. Tính Năng 01: PostGIS Deforestation Polygon & Copernicus Satellite Audit Engine")
add_body_p(doc_feat, "3. Tính Năng 02: Gemini 2.5 Flash OCR & Land Agreement Contract Auditor")
add_body_p(doc_feat, "4. Tính Năng 03: Mass Balance VICOFA Yield Ceiling Guard (3.500 kg/ha)")
add_body_p(doc_feat, "5. Tính Năng 04: 5-Year EUDR Passport & PostgreSQL Row-Level Security Vault")

add_heading_1(doc_feat, "1. TỔNG QUAN KIẾN TRÚC ĐỘNG CƠ TIỀN KIỂM")
add_body_p(doc_feat, "Coffee EU-Check AI đóng vai trò là 'Kiểm toán viên AI Nội bộ', tự động quét 100% tài liệu địa lý (GeoJSON), hợp đồng thuê đất nông hộ scan (Gemini 2.5 Flash OCR), kiểm soát trần sản lượng VICOFA (3.500 kg/ha) và cấp mã EUDR Passport lưu trữ 5 năm (EUDR Art. 31).")

# Feature 1
add_heading_1(doc_feat, "2. TÍNH NĂNG 01: POSTGIS POLYGON & COPERNICUS SATELLITE ENGINE")
add_heading_2(doc_feat, "· Tên tính năng: Động cơ Kiểm toán Tọa độ Địa lý PostGIS & Bản đồ Rừng Copernicus 2020")
add_body_p(doc_feat, "Nông hộ rẫy cà phê > 4.0 ha nộp 1 điểm GPS đơn lẻ (Point) thay vì chuỗi tọa độ đa giác Polygon (EUDR Art. 9(1)(d)). Hoặc đa giác bị chồng lấp với tọa độ rừng nguyên sinh theo bản đồ nền vệ tinh Copernicus 31/12/2020, dẫn đến rủi ro bị Hải quan EU giữ hàng và phạt 4% doanh thu.", "· Vấn đề nó giải: ")
add_body_p(doc_feat, "Tự động phân tích GeoJSON/KML qua PostGIS Spatial Engine & Turf.js. Đối soát diện tích rẫy thực tế với ranh giới bản đồ rừng vệ tinh Copernicus 2020. Hệ thống gắn cờ ERR_MISSING_POLYGON hoặc ERR_DEFORESTATION_OVERLAP dưới 10 giây.", "· Cách hoạt động: ")
add_body_p(doc_feat, "Loại bỏ 100% nguy cơ giữ container tại cảng EU do lỗi định dạng tọa độ; Rút ngắn thời gian kiểm tra ranh giới rẫy từ 14 ngày thủ công xuống < 10 giây.", "· Lợi ích đo được: ")

# Feature 2
add_heading_1(doc_feat, "3. TÍNH NĂNG 02: GEMINI 2.5 FLASH OCR & LAND CONTRACT AUDITOR")
add_heading_2(doc_feat, "· Tên tính năng: Động cơ Thẩm định Hợp đồng Đất Nông hộ Gemini 2.5 Flash OCR")
add_body_p(doc_feat, "Hàng ngàn hợp đồng thuê đất nông hộ viết tay hoặc scan nhòe đã hết hạn (EUDR Art. 3(b) yêu cầu đất hợp pháp). Kiểm tra thủ công gây trễ tiến độ nộp tờ khai và bỏ sót các hợp đồng hết hạn trước mốc giao hàng.", "· Vấn đề nó giải: ")
add_body_p(doc_feat, "Ứng dụng mô hình Gemini 2.5 Flash OCR tối ưu hóa cho tài liệu tiếng Việt. Tự động bóc tách ngày ký, thời hạn hiệu lực, tên chủ hộ, diện tích đất, và dán cờ đỏ ERR_EXPIRED_LAND_AGREEMENT nếu thời hạn hợp đồng không đủ bao phủ mốc xuất cảng.", "· Cách hoạt động: ")
add_body_p(doc_feat, "Tăng tốc độ xử lý file scan hợp đồng PDF/Image lên gấp 50 lần; Đạt độ chính xác bóc tách ngày tháng và diện tích 99,4%.", "· Lợi ích đo được: ")

# Feature 3
add_heading_1(doc_feat, "4. TÍNH NĂNG 03: MASS BALANCE VICOFA YIELD CEILING GUARD")
add_heading_2(doc_feat, "· Tên tính năng: Vòng Bảo vệ Trần Sản Lượng Cân Bằng Khối Lượng VICOFA (3.500 kg/ha)")
add_body_p(doc_feat, "Doanh nghiệp mua gom cà phê trôi nổi ngoài vùng trồng đăng ký. Khi sản lượng xuất khẩu khai báo vượt quá năng suất thực tế rẫy (trần VICOFA 3.500 kg/ha), Hải quan EU sẽ nghi vấn cà phê có nguồn gốc từ vùng phá rừng.", "· Vấn đề nó giải: ")
add_body_p(doc_feat, "Thuật toán Cân bằng Khối lượng (Mass Balance Guard) tính toán tổng sản lượng cà phê nhân xuất khẩu dựa trên diện tích rẫy hợp lệ. Nếu sản lượng/ha > 3.500 kg/ha, hệ thống lập tức chặn cấp Passport và yêu cầu bổ sung mã nông hộ đối soát.", "· Cách hoạt động: ")
add_body_p(doc_feat, "Đảm bảo 100% lô hàng xuất khẩu tuân thủ trần sản lượng VICOFA; Ngăn chặn triệt để rủi ro trà trộn cà phê không rõ nguồn gốc.", "· Lợi ích đo được: ")

# Feature 4
add_heading_1(doc_feat, "5. TÍNH NĂNG 04: 5-YEAR EUDR PASSPORT & POSTGRESQL RLS VAULT")
add_heading_2(doc_feat, "· Tên tính năng: Mã QR EUDR Passport 5 Năm & Hạ tầng Cách ly Dữ liệu PostgreSQL RLS")
add_body_p(doc_feat, "Nghĩa lưu trữ hồ sơ 5 năm theo EUDR Điều 31 và nỗi lo rò rỉ dữ liệu vùng thu mua thương mại độc quyền cho đối thủ cạnh tranh.", "· Vấn đề nó giải: ")
add_body_p(doc_feat, "Khởi tạo mã QR EUDR Passport chứa mã hóa băm (cryptographic hash) của toàn bộ hồ sơ tiền kiểm. Kết hợp cơ chế PostgreSQL Row-Level Security (RLS) cách ly tuyệt đối dữ liệu vùng thu mua giữa các doanh nghiệp xuất khẩu.", "· Cách hoạt động: ")
add_body_p(doc_feat, "Đáp ứng 100% nghĩa vụ lưu trữ 5 năm theo EUDR Art. 31; Bảo mật tuyệt đối 100% dữ liệu vùng trồng riêng tư của doanh nghiệp.", "· Lợi ích đo được: ")

doc_feat.save("docs/FEATURE_DOC_COFFEE_EU_CHECK_AI.docx")


# ==============================================================================
# 2. GENERATE PRICING DOC
# ==============================================================================
print("Generating Pricing Doc...")

doc_price = create_styled_doc(
    "TÀI LIỆU BẢNG GIÁ CƯỚC & MÔ HÌNH ROI (PRICING & ROI DOC)",
    "Coffee EU-Check AI — B2B SaaS Commercial Model & Return on Investment Analysis"
)

add_heading_1(doc_price, "1. 3 GÓI CƯỚC THƯƠNG MẠI & ĐỐI TƯỢNG PHÙ HỢP")

# Table for Pricing Tiers
table_p = doc_price.add_table(rows=4, cols=4)
table_p.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Tên Gói Cước", "Đơn Giá", "Đối Tượng Phù Hợp", "Quyền Lợi Cốt Lõi"]
for i, h in enumerate(headers):
    cell = table_p.cell(0, i)
    cell.text = h
    set_cell_background(cell, HEX_EMERALD)
    set_cell_margins(cell, 120, 120, 150, 150)
    p = cell.paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = COLOR_GOLD

tier_data = [
    ("Starter Pay-Per-Shipment", "1.200.000 VNĐ / lô hàng", "HTX & Doanh nghiệp vừa (5 - 15 container/năm)", "Quét 100% GeoJSON + OCR Hợp đồng + Cấp QR Passport 5 năm + Báo cáo Action Plan PDF"),
    ("Enterprise Professional (Khuyên Dùng)", "15.000.000 VNĐ / tháng", "Doanh nghiệp Xuất khẩu Top 20 (15 - 80 container/tháng)", "Không giới hạn số lô tiền kiểm + Động cơ PostGIS & Copernicus Vệ tinh + Hỗ trợ Hải quan EU 24/7 + RLS Dedicated Database"),
    ("Custom Infrastructure", "Báo giá B2B Tập đoàn Top 10", "Tập đoàn Xuất khẩu Đa quốc gia (> 100 container/tháng)", "Triển khai Private Cloud / On-Premise + Custom AI Model + API Webhook tích hợp ERP SAP/Oracle")
]

for row_idx, data in enumerate(tier_data, start=1):
    for col_idx, text in enumerate(data):
        cell = table_p.cell(row_idx, col_idx)
        cell.text = text
        set_cell_margins(cell, 100, 100, 120, 120)
        bg = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        set_cell_background(cell, bg)

doc_price.add_paragraph().paragraph_format.space_after = Pt(12)

add_heading_1(doc_price, "2. CÔNG THỨC ROI CHUẨN XUẤT KHẨU CÀ PHÊ EU")
add_heading_2(doc_price, "ĐẦU VÀO (INPUT)  --->  TIẾT KIỆM (SAVINGS)  --->  THỜI GIAN HOÀN VỐN (PAYBACK)")

add_heading_3(doc_price, "A. THÔNG SỐ ĐẦU VÀO (INPUT PARAMETERS)")
add_body_p(doc_price, "20 Container Cà Phê Robusta/Arabica xuất cảng EU per tháng.", "· Quy mô lô hàng: ")
add_body_p(doc_price, "$90.000 USD (Khoảng 2.25 Tỷ VNĐ) / Container 20ft.", "· Giá trị trung bình lô hàng: ")
add_body_p(doc_price, "$15.000 USD / Container (Tiền phạt lưu container, lưu bãi demurrage/detention và bồi thường hợp đồng đối tác EU).", "· Phí tổn rủi ro gián đoạn cảng EU: ")
add_body_p(doc_price, "4% Tổng doanh thu toàn cầu theo Điều 38 EUDR (Áp dụng khi bị kết luận cố tình xuất cà phê từ vùng phá rừng).", "· Áp lực phạt tối đa EUDR Art. 38: ")

add_heading_3(doc_price, "B. TÍNH TOÁN TIẾT KIỆM (MEASURABLE SAVINGS)")
add_body_p(doc_price, "Loại bỏ hoàn toàn nguy cơ 1 container bị giữ tại cảng EU do lỗi GeoJSON/Hợp đồng hết hạn -> Tiết kiệm trực tiếp $15.000 USD (375.000.000 VNĐ).", "1. Tránh chi phí phạt lưu cảng: ")
add_body_p(doc_price, "Giảm từ 3 nhân sự thẩm định thủ công (Chi phí 45.000.000 VNĐ/tháng) xuống 01 nhân sự vận hành AI SaaS (Tiết kiệm 30.000.000 VNĐ/tháng).", "2. Tiết kiệm chi phí nhân sự kiểm toán: ")
add_body_p(doc_price, "375.000.000 VNĐ (Phí lưu kho tránh được) + 30.000.000 VNĐ (Nhân sự) = 405.000.000 VNĐ / Tháng.", "3. TỔNG TIẾT KIỆM MỖI THÁNG: ")

add_heading_3(doc_price, "C. THỜI GIAN HOÀN VỐN (PAYBACK PERIOD)")
add_body_p(doc_price, "Chi phí gói Enterprise Pro: 15.000.000 VNĐ / Tháng.", "· Chi phí đầu tư SaaS: ")
add_body_p(doc_price, "Tỷ lệ ROI = (405.000.000 VNĐ Tiết kiệm / 15.000.000 VNĐ Chi phí SaaS) = 2700% (Gấp 27 lần).", "· Tỷ lệ ROI Hàng Tháng: ")
add_body_p(doc_price, "HOÀN VỐN NGAY TỪ LÔ HÀNG ĐẦU TIÊN (DƯỚI 03 NGÀY VẬN HÀNH).", "· THỜI GIAN HOÀN VỐN: ", )

doc_price.save("docs/PRICING_DOC_COFFEE_EU_CHECK_AI.docx")


# ==============================================================================
# 3. GENERATE SOLUTION DOC
# ==============================================================================
print("Generating Solution Doc...")

doc_sol = create_styled_doc(
    "TÀI LIỆU GIẢI PHÁP PHÂN KHÚC NGÀNH CÀ PHÊ (SOLUTION PLAYBOOK)",
    "Coffee EU-Check AI — Industry Segment Pain Points, Solutions & Real Case Studies"
)

# Segment 1
add_heading_1(doc_sol, "1. TẬP ĐOÀN XUẤT KHẨU CÀ PHÊ QUY MÔ LỚN (ENTERPRISE EXPORTERS)")
add_heading_2(doc_sol, "· Nỗi đau riêng (Pain Points):")
add_body_p(doc_sol, "Quản lý dữ liệu từ 20.000 đến 50.000 nông hộ liên kết tại Đắk Lắk và Gia Lai. Việc đối soát thủ công diện tích rẫy và hợp đồng thuê đất gây trễ tờ khai hải quan. Nguy cơ bị áp khung phạt 4% doanh thu toàn cầu theo EUDR Điều 38 cực kỳ nghiêm trọng.")
add_heading_2(doc_sol, "· Coffee EU-Check AI giải quyết thế nào:")
add_body_p(doc_sol, "Triển khai phân luồng dữ liệu tự động PostGIS Spatial Engine kết hợp PostgreSQL RLS Vault cách ly thông tin thu mua. Động cơ tự động quét 50.000 hồ sơ trong 15 phút, xuất báo cáo Action Plan PDF chuẩn Annex II cho từng container.")
add_heading_2(doc_sol, "· Câu chuyện mẫu (Case Study Story):")
add_body_p(doc_sol, "Tập đoàn Xuất khẩu Cà phê X (Buôn Ma Thuột) trước kỳ hạn EUDR chuẩn bị lô hàng 40 container sang Cảng Rotterdam. Qua kiểm toán thử nghiệm, Coffee EU-Check AI phát hiện 12% tọa độ rẫy bị trôi vào ranh giới rừng Copernicus 2020 và 350 hợp đồng đất nông hộ đã hết hạn. Hệ thống đưa ra cảnh báo đỏ, giúp Tập đoàn X kịp thời điều chỉnh vùng thu mua thay thế, loại bỏ nguy cơ bị giữ kho đền bù $600.000 USD tại Cảng Rotterdam.")

# Segment 2
add_heading_1(doc_sol, "2. ĐƠN VỊ THƯƠNG MẠI TRUNG GIAN & TRADING HOUSES")
add_heading_2(doc_sol, "· Nỗi đau riêng (Pain Points):")
add_body_p(doc_sol, "Mua gom cà phê nhân xô từ hàng trăm đại lý thu mua nhỏ lẻ. Cà phê trôi nổi không rõ nguồn gốc dễ vượt trần sản lượng 3.500 kg/ha của VICOFA, dẫn đến việc nhà xuất khẩu lớn từ chối thu mua do không chứng minh được tính hợp pháp.")
add_heading_2(doc_sol, "· Coffee EU-Check AI giải quyết thế nào:")
add_body_p(doc_sol, "Cung cấp công cụ Mass Balance Yield Ceiling Guard. Thuật toán tự động siết trần 3.500 kg/ha trên từng mã nông hộ, tự động lọc bỏ các mã cà phê trôi nổi và cấp chứng nhận tiền kiểm giúp Trading House tự tin bán cho các Tập đoàn xuất khẩu lớn.")
add_heading_2(doc_sol, "· Câu chuyện mẫu (Case Study Story):")
add_body_p(doc_sol, "Trading House Y tại Gia Lai thu mua 1.500 tấn cà phê xô. Nhờ sử dụng gói Starter Pay-Per-Shipment của Coffee EU-Check AI, Trading House Y đã tự động cấp mã QR EUDR Passport cho 100% lô hàng. Nhờ hồ sơ minh bạch đạt Readiness Score 98/100, lô hàng của Trading House Y được các tập đoàn xuất khẩu thu mua với giá cao hơn 2% so với giá thị trường.")

# Segment 3
add_heading_1(doc_sol, "3. HỢP TÁC XÃ & LIÊN MINH NÔNG HỘ TÂY NGUYÊN")
add_heading_2(doc_sol, "· Nỗi đau riêng (Pain Points):")
add_body_p(doc_sol, "Thiếu nhân sự chuyên trách kỹ thuật GIS và luật EU. Nông dân chỉ chấm 1 điểm GPS tại nhà thay vì vẽ Polygon rẫy cà phê, dẫn đến rủi ro toàn bộ sản lượng của HTX bị coi là vi phạm EUDR Art. 9(1)(d).")
add_heading_2(doc_sol, "· Coffee EU-Check AI giải quyết thế nào:")
add_body_p(doc_sol, "Giao diện đơn giản hóa cho Ban Quản trị HTX. Tích hợp công cụ chuyển đổi tự động từ file KML/GPX điện thoại thành đa giác Polygon chuẩn GeoJSON, kết hợp Gemini 2.5 Flash OCR đọc hợp đồng scan chụp bằng smartphone.")
add_heading_2(doc_sol, "· Câu chuyện mẫu (Case Study Story):")
add_body_p(doc_sol, "HTX Cà phê Bền vững Z (Lâm Đồng) gồm 180 hộ thành viên. Ban Quản trị HTX đã dùng smartphone chụp 180 hợp đồng đất và vẽ ranh giới rẫy qua app. Coffee EU-Check AI xử lý và trả về 180 mã EUDR Passport chuẩn hóa chỉ sau 1 giờ, giúp HTX ký kết thành công hợp đồng cung ứng 300 tấn cà phê đạt chuẩn EUDR cho nhà xuất khẩu.")

# Segment 4
add_heading_1(doc_sol, "4. NHÀ RANG XAY & XUẤT KHẨU CHUYÊN BIỆT (SPECIALTY ROASTERS)")
add_heading_2(doc_sol, "· Nỗi đau riêng (Pain Points):")
add_body_p(doc_sol, "Xuất khẩu các dòng cà phê Fine Coffee / Specialty chất lượng cao sang các thị trường khó tính như Đức, Hà Lan. Cần minh chứng tính bền vững khắt khe để giữ uy tín thương hiệu cao cấp và nỗi lo rò rỉ dữ liệu nguồn gốc cho đối thủ.")
add_heading_2(doc_sol, "· Coffee EU-Check AI giải quyết thế nào:")
add_body_p(doc_sol, "Cấp thẻ EUDR Passport gắn trực tiếp lên từng bao bì lô hàng xuất khẩu với hạ tầng PostgreSQL RLS bảo vệ dữ liệu thương mại riêng tư tuyệt đối. Đối tác EU quét mã QR có thể xem ngay chứng nhận tiền kiểm mà không thấy dữ liệu thương mại nhạy cảm.")
add_heading_2(doc_sol, "· Câu chuyện mẫu (Case Study Story):")
add_body_p(doc_sol, "Thương hiệu Cà phê Chuyên biệt W xuất khẩu container cà phê Fine Robusta sang Frankfurt. Nhờ gắn mã EUDR Passport từ Coffee EU-Check AI, lô hàng được Hải quan Đức thông quan luồng xanh tức thì mà không cần giữ kho kiểm tra, giúp giữ nguyên hương vị cà phê tươi mới và khẳng định vị thế thương hiệu cao cấp.")

doc_sol.save("docs/SOLUTION_DOC_COFFEE_EU_CHECK_AI.docx")

# ==============================================================================
# 4. GENERATE COMPANION MARKDOWN FILES IN Content/
# ==============================================================================
print("Generating Companion Markdown Files in Content/...")

with open("Content/FEATURE_DOC.md", "w", encoding="utf-8") as f:
    f.write("""# TÀI LIỆU ĐỘNG CƠ TÍNH NĂNG TIỀN KIỂM (FEATURE DOC)
**Product:** Coffee EU-Check AI — Themis Regal Legal-Tech Pre-Audit System

## MỤC LỤC TỰ ĐỘNG
1. Tổng Quan Kiến Trúc Động Cơ Kép (Dual GIS & AI Engine)
2. Tính Năng 01: PostGIS Deforestation Polygon & Copernicus Satellite Audit Engine
3. Tính Năng 02: Gemini 2.5 Flash OCR & Land Agreement Contract Auditor
4. Tính Năng 03: Mass Balance VICOFA Yield Ceiling Guard (3.500 kg/ha)
5. Tính Năng 04: 5-Year EUDR Passport & PostgreSQL Row-Level Security Vault

---

## 1. TỔNG QUAN KIẾN TRÚC
Coffee EU-Check AI đóng vai trò là 'Kiểm toán viên AI Nội bộ', tự động quét 100% tài liệu địa lý (GeoJSON), hợp đồng thuê đất nông hộ scan (Gemini 2.5 Flash OCR), kiểm soát trần sản lượng VICOFA (3.500 kg/ha) và cấp mã EUDR Passport lưu trữ 5 năm (EUDR Art. 31).

---

## 2. TÍNH NĂNG 01: POSTGIS POLYGON & COPERNICUS SATELLITE ENGINE
- **Tên tính năng:** Động cơ Kiểm toán Tọa độ Địa lý PostGIS & Bản đồ Rừng Copernicus 2020
- **Vấn đề nó giải:** Nông hộ rẫy cà phê > 4.0 ha nộp 1 điểm GPS đơn lẻ (Point) thay vì chuỗi tọa độ đa giác Polygon (EUDR Art. 9(1)(d)). Hoặc đa giác bị chồng lấp với tọa độ rừng nguyên sinh theo bản đồ nền vệ tinh Copernicus 31/12/2020.
- **Cách hoạt động:** Tự động phân tích GeoJSON/KML qua PostGIS Spatial Engine & Turf.js. Đối soát diện tích rẫy thực tế với ranh giới bản đồ rừng vệ tinh Copernicus 2020. Hệ thống gắn cờ ERR_MISSING_POLYGON hoặc ERR_DEFORESTATION_OVERLAP dưới 10 giây.
- **Lợi ích đo được:** Loại bỏ 100% nguy cơ giữ container tại cảng EU do lỗi định dạng tọa độ; Rút ngắn thời gian kiểm tra ranh giới rẫy từ 14 ngày thủ công xuống < 10 giây.

---

## 3. TÍNH NĂNG 02: GEMINI 2.5 FLASH OCR & LAND CONTRACT AUDITOR
- **Tên tính năng:** Động cơ Thẩm định Hợp đồng Đất Nông hộ Gemini 2.5 Flash OCR
- **Vấn đề nó giải:** Hàng ngàn hợp đồng thuê đất nông hộ viết tay hoặc scan nhòe đã hết hạn (EUDR Art. 3(b) yêu cầu đất hợp pháp). Kiểm tra thủ công gây trễ tiến độ nộp tờ khai.
- **Cách hoạt động:** Ứng dụng mô hình Gemini 2.5 Flash OCR tối ưu hóa cho tài liệu tiếng Việt. Tự động bóc tách ngày ký, thời hạn hiệu lực, tên chủ hộ, diện tích đất, và dán cờ đỏ ERR_EXPIRED_LAND_AGREEMENT nếu hợp đồng hết hạn.
- **Lợi ích đo được:** Tăng tốc độ xử lý file scan hợp đồng PDF/Image lên gấp 50 lần; Đạt độ chính xác bóc tách ngày tháng và diện tích 99,4%.

---

## 4. TÍNH NĂNG 03: MASS BALANCE VICOFA YIELD CEILING GUARD
- **Tên tính năng:** Vòng Bảo vệ Trần Sản Lượng Cân Bằng Khối Lượng VICOFA (3.500 kg/ha)
- **Vấn đề nó giải:** Doanh nghiệp mua gom cà phê trôi nổi ngoài vùng trồng đăng ký. Khi sản lượng xuất khẩu khai báo vượt quá năng suất thực tế rẫy (trần VICOFA 3.500 kg/ha), Hải quan EU sẽ nghi vấn cà phê có nguồn gốc từ vùng phá rừng.
- **Cách hoạt động:** Thuật toán Cân bằng Khối lượng (Mass Balance Guard) tính toán tổng sản lượng cà phê nhân xuất khẩu dựa trên diện tích rẫy hợp lệ. Nếu sản lượng/ha > 3.500 kg/ha, hệ thống lập tức chặn cấp Passport.
- **Lợi ích đo được:** Đảm bảo 100% lô hàng xuất khẩu tuân thủ trần sản lượng VICOFA; Ngăn chặn triệt để rủi ro trà trộn cà phê không rõ nguồn gốc.

---

## 5. TÍNH NĂNG 04: 5-YEAR EUDR PASSPORT & POSTGRESQL RLS VAULT
- **Tên tính năng:** Mã QR EUDR Passport 5 Năm & Hạ tầng Cách ly Dữ liệu PostgreSQL RLS
- **Vấn đề nó giải:** Nghĩa vụ lưu trữ hồ sơ 5 năm theo EUDR Điều 31 và nỗi lo rò rỉ dữ liệu vùng thu mua thương mại độc quyền cho đối thủ cạnh tranh.
- **Cách hoạt động:** Khởi tạo mã QR EUDR Passport chứa mã hóa băm của toàn bộ hồ sơ tiền kiểm. Kết hợp cơ chế PostgreSQL Row-Level Security (RLS) cách ly tuyệt đối dữ liệu vùng thu mua.
- **Lợi ích đo được:** Đáp ứng 100% nghĩa vụ lưu trữ 5 năm theo EUDR Art. 31; Bảo mật tuyệt đối 100% dữ liệu vùng trồng riêng tư của doanh nghiệp.
""")

with open("Content/PRICING_DOC.md", "w", encoding="utf-8") as f:
    f.write("""# TÀI LIỆU BẢNG GIÁ CƯỚC & MÔ HÌNH ROI (PRICING & ROI DOC)
**Product:** Coffee EU-Check AI — B2B SaaS Commercial Model & Return on Investment

## 1. 3 GÓI CƯỚC THƯƠNG MẠI & ĐỐI TƯỢNG PHÙ HỢP

| Tên Gói Cước | Đơn Giá | Đối Tượng Phù Hợp | Quyền Lợi Cốt Lõi |
| :--- | :--- | :--- | :--- |
| **Starter Pay-Per-Shipment** | 1.200.000 VNĐ / lô hàng | HTX & Doanh nghiệp vừa (5 - 15 container/năm) | Quét 100% GeoJSON + OCR Hợp đồng + Cấp QR Passport 5 năm + Báo cáo PDF |
| **Enterprise Professional (Khuyên Dùng)** | 15.000.000 VNĐ / tháng | Doanh nghiệp Xuất khẩu Top 20 (15 - 80 container/tháng) | Không giới hạn lô + PostGIS & Copernicus Vệ tinh + Hỗ trợ Hải quan EU 24/7 + RLS Database |
| **Custom Infrastructure** | Báo giá B2B Tập đoàn Top 10 | Tập đoàn Xuất khẩu Đa quốc gia (> 100 container/tháng) | Private Cloud / On-Premise + Custom AI Model + API Webhook tích hợp ERP SAP/Oracle |

---

## 2. CÔNG THỨC ROI CHUẨN XUẤT KHẨU CÀ PHÊ EU
`ĐẦU VÀO (INPUT) ---> TIẾT KIỆM (SAVINGS) ---> THỜI GIAN HOÀN VỐN (PAYBACK)`

### A. Thông số đầu vào (Input)
- **Quy mô lô hàng:** 20 Container Cà Phê Robusta/Arabica xuất cảng EU / tháng.
- **Giá trị trung bình lô hàng:** $90.000 USD (Khoảng 2.25 Tỷ VNĐ) / Container 20ft.
- **Phí tổn rủi ro lưu cảng:** $15.000 USD / Container (Lưu container, lưu bãi đền bù).
- **Áp lực phạt tối đa:** 4% Tổng doanh thu toàn cầu theo Điều 38 EUDR.

### B. Tính toán tiết kiệm (Savings)
1. **Tránh phí lưu cảng:** Loại bỏ nguy cơ 1 container bị giữ -> Tiết kiệm $15.000 USD (375.000.000 VNĐ).
2. **Tiết kiệm nhân sự:** Giảm từ 3 nhân sự thẩm định (45M đ/tháng) xuống 1 nhân sự AI SaaS (15M đ/tháng) -> Tiết kiệm 30.000.000 VNĐ/tháng.
3. **TỔNG TIẾT KIỆM HÀNG THÁNG:** 405.000.000 VNĐ / Tháng.

### C. Thời gian hoàn vốn (Payback Period)
- **Chi phí SaaS:** 15.000.000 VNĐ / Tháng (Gói Enterprise Pro).
- **Tỷ lệ ROI:** 405.000.000 / 15.000.000 = **2700% (Gấp 27 lần)**.
- **THỜI GIAN HOÀN VỐN:** **HOÀN VỐN NGAY TỪ LÔ HÀNG ĐẦU TIÊN (DƯỚI 03 NGÀY VẬN HÀNH)**.
""")

with open("Content/SOLUTION_DOC.md", "w", encoding="utf-8") as f:
    f.write("""# TÀI LIỆU GIẢI PHÁP PHÂN KHÚC NGÀNH CÀ PHÊ (SOLUTION PLAYBOOK)
**Product:** Coffee EU-Check AI — Industry Segment Solution & Case Studies

## 1. TẬP ĐOÀN XUẤT KHẨU CÀ PHÊ QUY MÔ LỚN (ENTERPRISE EXPORTERS)
- **Nỗi đau riêng:** Quản lý 20.000 - 50.000 nông hộ tại Đắk Lắk/Gia Lai. Đối soát diện tích rẫy và hợp đồng thuê đất thủ công gây trễ tờ khai. Nguy cơ bị phạt 4% doanh thu toàn cầu theo EUDR Art. 38.
- **Coffee EU-Check AI giải quyết thế nào:** Phân luồng dữ liệu tự động PostGIS Spatial Engine & PostgreSQL RLS Vault. Quét 50.000 hồ sơ trong 15 phút, xuất báo cáo Action Plan PDF chuẩn Annex II.
- **Câu chuyện mẫu:** Tập đoàn X (Buôn Ma Thuột) xuất 40 container sang Cảng Rotterdam. Coffee EU-Check AI phát hiện 12% tọa độ rẫy bị trôi vào ranh giới rừng Copernicus 2020 và 350 hợp đồng đất hết hạn, giúp Tập đoàn X kịp thời điều chỉnh vùng thu mua, tránh phí đền bù $600.000 USD tại Cảng Rotterdam.

---

## 2. ĐƠN VỊ THƯƠNG MẠI TRUNG GIAN & TRADING HOUSES
- **Nỗi đau riêng:** Mua gom cà phê từ hàng trăm đại lý nhỏ lẻ. Cà phê trôi nổi dễ vượt trần 3.500 kg/ha VICOFA, dẫn đến việc nhà xuất khẩu lớn từ chối thu mua do không chứng minh được nguồn gốc.
- **Coffee EU-Check AI giải quyết thế nào:** Thuật toán Mass Balance Yield Ceiling Guard tự động siết trần 3.500 kg/ha trên từng mã nông hộ, tự động lọc bỏ cà phê trôi nổi và cấp chứng nhận tiền kiểm.
- **Câu chuyện mẫu:** Trading House Y tại Gia Lai thu mua 1.500 tấn cà phê. Dùng gói Starter Pay-Per-Shipment cấp mã QR EUDR Passport cho 100% lô hàng. Nhờ Readiness Score 98/100, lô hàng được các tập đoàn lớn thu mua với giá cao hơn 2% so với giá thị trường.

---

## 3. HỢP TÁC XÃ & LIÊN MINH NÔNG HỘ TÂY NGUYÊN
- **Nỗi đau riêng:** Thiếu nhân sự chuyên trách kỹ thuật GIS và luật EU. Nông dân chỉ chấm 1 điểm GPS tại nhà thay vì vẽ Polygon rẫy cà phê, vi phạm EUDR Art. 9(1)(d).
- **Coffee EU-Check AI giải quyết thế nào:** Giao diện đơn giản hóa cho Ban Quản trị HTX. Tích hợp công cụ chuyển đổi KML/GPX từ điện thoại thành Polygon chuẩn GeoJSON và Gemini 2.5 Flash OCR đọc hợp đồng scan chụp bằng smartphone.
- **Câu chuyện mẫu:** HTX Cà phê Bền vững Z (Lâm Đồng) gồm 180 hộ thành viên. Dùng smartphone chụp 180 hợp đồng đất và vẽ ranh giới rẫy. Coffee EU-Check AI trả về 180 mã EUDR Passport chuẩn hóa sau 1 giờ, giúp HTX ký thành công hợp đồng cung ứng 300 tấn cà phê đạt chuẩn EUDR.

---

## 4. NHÀ RANG XAY & XUẤT KHẨU CHUYÊN BIỆT (SPECIALTY ROASTERS)
- **Nỗi đau riêng:** Xuất khẩu Fine Coffee / Specialty sang Đức, Hà Lan. Cần minh chứng tính bền vững khắt khe để giữ uy tín thương hiệu cao cấp và nỗi lo rò rỉ dữ liệu nguồn gốc cho đối thủ.
- **Coffee EU-Check AI giải quyết thế nào:** Cấp thẻ EUDR Passport gắn trực tiếp lên từng bao bì lô hàng với hạ tầng PostgreSQL RLS bảo vệ dữ liệu thương mại riêng tư tuyệt đối.
- **Câu chuyện mẫu:** Thương hiệu Cà phê Chuyên biệt W xuất khẩu container Fine Robusta sang Frankfurt. Gắn mã EUDR Passport, lô hàng được Hải quan Đức thông quan luồng xanh tức thì mà không cần giữ kho kiểm tra, giữ nguyên hương vị tươi mới và khẳng định uy tín thương hiệu cao cấp.
""")

print("ALL DOCUMENTS GENERATED SUCCESSFULLY!")
